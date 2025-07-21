from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Request, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
import os
import json
import uuid
from datetime import datetime
from typing import List, Dict, Any
import aiofiles
import mammoth
import io

# Import from parent directory (backend)
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from celery_worker import process_audio_task, generate_summary_task
from utils import (
    find_input_audio, get_processed_dir, get_transcript_path, 
    get_summary_path, get_metadata_path, update_metadata
)
from prompt_manager import get_prompt_manager, format_prompt, list_prompts, reload_prompts
from timing_model import timing_model

app = FastAPI(title="Audio Transcription & Analysis API", version="1.0.0")

# CORS middleware for React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # React dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files (only if directory exists)
if os.path.exists("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")

def resolve_path_from_metadata(path: str) -> str:
    """Resolve file paths from metadata to work with backend directory structure"""
    if not path:
        return path
    
    # If it's already an absolute path, return as is
    if os.path.isabs(path):
        return path
    
    # If it starts with 'data/', make it relative to project root (parent of backend)
    if path.startswith('data/'):
        return os.path.join("..", path)
    
    # If it already starts with '..', it's already resolved correctly
    if path.startswith('..'):
        return path
    
    # If it's a relative path, assume it's relative to project root
    return os.path.join("..", path)

def get_audio_metadata(audio_id: str) -> Dict[str, Any]:
    """Get metadata for a specific audio ID"""
    try:
        # Update path to look in parent directory for data
        metadata_dir = os.path.join("..", "data", audio_id, "metadata")
        print(f"[DEBUG] Trying to access metadata directory: {metadata_dir}")
        if not os.path.exists(metadata_dir):
            print(f"[ERROR] Metadata directory does not exist: {metadata_dir}")
            return None
        
        metadata_files = [f for f in os.listdir(metadata_dir) if f.endswith('.json')]
        if not metadata_files:
            print(f"[ERROR] No metadata JSON files found in: {metadata_dir}")
            return None
        
        metadata_path = os.path.join(metadata_dir, metadata_files[0])
        print(f"[DEBUG] Reading metadata file: {metadata_path}")
        with open(metadata_path, 'r', encoding='utf-8') as f:
            metadata = json.load(f)
            
        # Fix paths in metadata to work with backend directory structure
        if 'transcript_path' in metadata:
            metadata['transcript_path'] = resolve_path_from_metadata(metadata['transcript_path'])
        if 'summary_path' in metadata:
            metadata['summary_path'] = resolve_path_from_metadata(metadata['summary_path'])
        if 'audio_path' in metadata:
            metadata['audio_path'] = resolve_path_from_metadata(metadata['audio_path'])
            
        return metadata
    except Exception as e:
        print(f"[ERROR] Exception reading metadata for {audio_id}: {e}")
        return None

def get_all_audio_metadata() -> List[Dict[str, Any]]:
    """Get metadata for all audio files"""
    audios = []
    print("[DEBUG] Checking if data directory exists...")
    # Update path to look in parent directory for data
    data_dir = os.path.join("..", "data")
    if not os.path.exists(data_dir):
        print("[ERROR] Data directory does not exist!")
        return audios
    print("[DEBUG] Listing audio IDs in data directory...")
    
    # Get all audio directories and sort them by audio_id (numeric)
    audio_dirs = []
    for audio_id in os.listdir(data_dir):
        audio_dir = os.path.join(data_dir, audio_id)
        if os.path.isdir(audio_dir) and audio_id.isdigit():
            audio_dirs.append((int(audio_id), audio_dir))
    
    # Sort by audio_id in descending order (newest first)
    audio_dirs.sort(key=lambda x: x[0], reverse=True)
    
    for audio_id, audio_dir in audio_dirs:
        print(f"[DEBUG] Checking audio directory: {audio_dir}")
        metadata = get_audio_metadata(str(audio_id))
        if metadata:
            metadata['audio_id'] = str(audio_id)
            audios.append(metadata)
        else:
            print(f"[WARN] No metadata found for audio_id: {audio_id}")
    
    return audios

@app.get("/")
async def root():
    return {"message": "Audio Transcription & Analysis API"}

@app.post("/upload")
async def upload_audio(
    file: UploadFile = File(...),
    speedup: float = Form(1.0),
    auto_adjust: bool = Form(False),
    chunk: bool = Form(False),
    chunk_duration: int = Form(10),
    diarizer: str = Form("huggingface")
):
    """Upload audio file and start processing"""
    
    # Validate file type
    allowed_extensions = {'.wav', '.mp3', '.m4a', '.flac', '.aac'}
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Invalid file type")
    
    # Generate sequential audio ID starting from 1000
    def get_next_audio_id():
        data_dir = os.path.join("..", "data")
        if not os.path.exists(data_dir):
            return "1000"
        
        existing_ids = []
        for item in os.listdir(data_dir):
            if os.path.isdir(os.path.join(data_dir, item)) and item.isdigit():
                existing_ids.append(int(item))
        
        if not existing_ids:
            return "1000"
        
        next_id = max(existing_ids) + 1
        return str(next_id)
    
    audio_id = get_next_audio_id()
    
    # Create directory structure - update path to parent directory
    input_dir = os.path.join("..", "data", audio_id, "input_audio")
    os.makedirs(input_dir, exist_ok=True)
    
    # Save uploaded file
    file_path = os.path.join(input_dir, file.filename)
    async with aiofiles.open(file_path, 'wb') as f:
        content = await file.read()
        await f.write(content)
    
    # Get actual audio duration immediately after upload
    try:
        from audio_processor import get_audio_duration
        actual_duration_seconds = get_audio_duration(file_path)
        actual_duration_minutes = actual_duration_seconds / 60.0
        print(f"[DEBUG] Detected audio duration at upload: {actual_duration_minutes:.2f} minutes")
    except Exception as e:
        print(f"[WARNING] Could not determine audio duration at upload: {e}")
        actual_duration_minutes = None
    
    # Debug: Log the received parameters
    print(f"[DEBUG] Upload parameters received:")
    print(f"  - speedup: {speedup}")
    print(f"  - auto_adjust: {auto_adjust}")
    print(f"  - chunk: {chunk}")
    print(f"  - chunk_duration: {chunk_duration}")
    print(f"  - diarizer: {diarizer}")
    print(f"  - actual_duration_minutes: {actual_duration_minutes}")
    
    # Start background processing
    task = process_audio_task.delay(
        audio_id=audio_id,
        filename=file.filename,
        speedup=speedup,
        auto_adjust=auto_adjust,
        chunk=chunk,
        chunk_duration=chunk_duration,
        diarizer=diarizer,
        actual_duration_minutes=actual_duration_minutes  # Pass actual duration to task
    )
    
    return {
        "audio_id": audio_id,
        "filename": file.filename,
        "task_id": task.id,
        "status": "uploaded",
        "message": "Audio uploaded and processing started"
    }

@app.get("/status/{audio_id}")
async def get_audio_status(audio_id: str):
    """Get processing status for a specific audio"""
    metadata = get_audio_metadata(audio_id)
    if not metadata:
        raise HTTPException(status_code=404, detail="Audio not found")
    
    return metadata

@app.get("/transcript/{audio_id}")
async def get_transcript(audio_id: str):
    """Get transcript for a specific audio"""
    metadata = get_audio_metadata(audio_id)
    if not metadata:
        raise HTTPException(status_code=404, detail="Audio not found")
    
    if metadata.get("status") not in ["transcribed", "summary_generated", "summary_regenerating"]:
        raise HTTPException(status_code=400, detail="Transcript not ready yet")
    
    transcript_path = metadata.get("transcript_path")
    if not transcript_path:
        raise HTTPException(status_code=404, detail="Transcript file not found")
    
    # Path is already resolved in get_audio_metadata
    if not os.path.exists(transcript_path):
        raise HTTPException(status_code=404, detail="Transcript file not found")
    
    return FileResponse(transcript_path, media_type="text/csv")



@app.get("/document/{audio_id}")
async def get_document(audio_id: str):
    """Get document/summary for a specific audio"""
    metadata = get_audio_metadata(audio_id)
    if not metadata:
        raise HTTPException(status_code=404, detail="Audio not found")
    
    if metadata.get("status") != "summary_generated":
        raise HTTPException(status_code=400, detail="Document not ready yet")
    
    summary_path = metadata.get("summary_path")
    if not summary_path:
        raise HTTPException(status_code=404, detail="Document file not found")
    
    # Path is already resolved in get_audio_metadata
    if not os.path.exists(summary_path):
        raise HTTPException(status_code=404, detail="Document file not found")
    
    return FileResponse(summary_path, media_type="text/plain")

def _add_formatted_content_to_doc(doc, content):
    """
    Parse markdown content and add it to Word document with formatting that matches 
    the ReactMarkdown styling used in the frontend UI.
    """
    import re
    
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines but add spacing
        if not line:
            i += 1
            continue
        
        # Skip table separator lines completely (---, ===, |---|)
        if re.match(r'^[\|\s]*[-=]+[\|\s]*$', line):
            i += 1
            continue
        
        # Handle tables
        if '|' in line and not re.match(r'^[\|\s]*[-=]+[\|\s]*$', line):
            table_lines = []
            j = i
            while j < len(lines):
                current_line = lines[j].strip()
                if '|' in current_line:
                    # Skip separator lines in table
                    if not re.match(r'^[\|\s]*[-=]+[\|\s]*$', current_line):
                        table_lines.append(current_line)
                    j += 1
                else:
                    break
            
            if len(table_lines) >= 1:  # At least one actual row
                _add_table_to_doc(doc, table_lines)
                i = j
                continue
        
        # Handle headers (# ## ### ####) - matching ReactMarkdown h1-h4
        if line.startswith('#'):
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1
            
            header_text = line[level:].strip()
            if header_text:
                _add_header_to_doc(doc, header_text, level)
            i += 1
            continue
        
        # Handle indented sub-items first (more specific pattern)
        if re.match(r'^\s{2,}[-*•]\s+', line):
            indent_level = len(line) - len(line.lstrip())
            bullet_text = re.sub(r'^\s+[-*•]\s+', '', line)
            
            # Create indented bullet point
            paragraph = doc.add_paragraph()
            paragraph.style = 'List Bullet'
            # Add proper indentation for sub-items
            paragraph.paragraph_format.left_indent = 720000  # 0.5 inch indent
            paragraph.paragraph_format.first_line_indent = -360000  # Hanging indent
            _add_formatted_text_to_paragraph(paragraph, bullet_text)
            paragraph.paragraph_format.space_after = 50000
            i += 1
            continue
        
        # Handle main bullet points (- or * or •) - matching ReactMarkdown ul
        if line.startswith('- ') or line.startswith('* ') or line.startswith('• '):
            bullet_text = line[2:].strip()
            paragraph = doc.add_paragraph(style='List Bullet')
            _add_formatted_text_to_paragraph(paragraph, bullet_text)
            # Add slight spacing like ReactMarkdown (mb-1)
            paragraph.paragraph_format.space_after = 50000  # 0.05 inches
            i += 1
            continue
        
        # Handle numbered lists (1. 2. etc) - matching ReactMarkdown ol
        if re.match(r'^\d+\.\s+', line):
            list_text = re.sub(r'^\d+\.\s+', '', line)
            paragraph = doc.add_paragraph(style='List Number')
            _add_formatted_text_to_paragraph(paragraph, list_text)
            # Add slight spacing like ReactMarkdown (mb-1)
            paragraph.paragraph_format.space_after = 50000  # 0.05 inches
            i += 1
            continue
        
        # Handle regular paragraphs - matching ReactMarkdown p (mb-3 leading-relaxed)
        if line:
            paragraph = doc.add_paragraph()
            _add_formatted_text_to_paragraph(paragraph, line)
            # Add spacing like ReactMarkdown (mb-3)
            paragraph.paragraph_format.space_after = 150000  # 0.15 inches
            # Set line spacing like ReactMarkdown (leading-relaxed)
            paragraph.paragraph_format.line_spacing = 1.3
        
        i += 1

def _add_header_to_doc(doc, text, level):
    """Add header with styling that matches ReactMarkdown h1-h4"""
    from docx.shared import Pt
    
    # Match ReactMarkdown header levels and styling
    if level == 1:
        # h1: text-2xl font-bold (24px, bold)
        heading = doc.add_heading(text, level=1)
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.bold = True
        # mt-6 mb-4 spacing
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(12)
        
    elif level == 2:
        # h2: text-xl font-semibold (20px, semibold)
        heading = doc.add_heading(text, level=2)
        heading.runs[0].font.size = Pt(20)
        heading.runs[0].font.bold = True
        # mt-5 mb-3 spacing
        heading.paragraph_format.space_before = Pt(15)
        heading.paragraph_format.space_after = Pt(9)
        
    elif level == 3:
        # h3: text-lg font-medium (18px, medium)
        heading = doc.add_heading(text, level=3)
        heading.runs[0].font.size = Pt(18)
        heading.runs[0].font.bold = True
        # mt-4 mb-2 spacing
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        
    elif level == 4:
        # h4: text-base font-medium (16px, medium)
        heading = doc.add_heading(text, level=4)
        heading.runs[0].font.size = Pt(16)
        heading.runs[0].font.bold = True
        # mt-3 mb-2 spacing
        heading.paragraph_format.space_before = Pt(9)
        heading.paragraph_format.space_after = Pt(6)
        
    else:
        # Default for h5+ levels
        heading = doc.add_heading(text, level=min(level, 9))
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True

def _add_table_to_doc(doc, table_lines):
    """Add table with styling that matches ReactMarkdown table formatting"""
    from docx.shared import Pt
    import re
    
    # Parse table structure and completely exclude separator lines
    rows = []
    for line in table_lines:
        line = line.strip()
        # Skip any line that's primarily dashes, equals, or just pipes and spaces
        if re.match(r'^[\|\s]*[-=]+[\|\s]*$', line) or not line:
            continue
        
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first/last cells (common in markdown tables)
        if cells and cells[0] == '':
            cells = cells[1:]
        if cells and cells[-1] == '':
            cells = cells[:-1]
        
        if cells:  # Only add rows with actual content
            rows.append(cells)
    
    if not rows:
        return
    
    # Create Word table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    
    # Format table to match ReactMarkdown styling
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            if j < len(table.rows[i].cells):  # Ensure cell exists
                cell = table.rows[i].cells[j]
                cell.text = cell_data
                
                # Header row styling (thead bg-gray-50, font-medium)
                if i == 0:
                    cell.paragraphs[0].runs[0].font.bold = True
                    # Set background color (light gray like bg-gray-50)
                    from docx.oxml.shared import qn
                    from docx.oxml import parse_xml
                    shading_elm = parse_xml(r'<w:shd {} w:fill="F9FAFB"/>'.format(
                        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # Add padding like ReactMarkdown (px-3 py-2)
                cell.paragraphs[0].paragraph_format.left_indent = Pt(9)
                cell.paragraphs[0].paragraph_format.right_indent = Pt(9)
    
    # Add spacing like ReactMarkdown (my-4)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)

def _process_inline_formatting(text):
    """Process inline markdown formatting like **bold** and *italic*"""
    return text

def _add_formatted_text_to_paragraph(paragraph, text):
    """Add text to paragraph with markdown formatting converted to Word formatting"""
    import re
    
    # Handle bold text **text** - matching ReactMarkdown strong (font-semibold)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text (font-semibold)
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text - matching ReactMarkdown em (italic)
            italic_text = part[1:-1]
            run = paragraph.add_run(italic_text)
            run.italic = True
        else:
            # Regular text
            if part:
                paragraph.add_run(part)

@app.get("/export-document/{audio_id}")
async def export_document_word(audio_id: str):
    """Export document as Word file"""
    try:
        print(f"[DEBUG] Export Word document requested for audio_id: {audio_id}")
        
        # Step 1: Get metadata
        metadata = get_audio_metadata(audio_id)
        if not metadata:
            print(f"[ERROR] No metadata found for audio_id: {audio_id}")
            raise HTTPException(status_code=404, detail="Audio not found")
        
        print(f"[DEBUG] Metadata found. Status: {metadata.get('status')}")
        
        # Step 2: Check status
        if metadata.get("status") != "summary_generated":
            print(f"[ERROR] Document not ready. Current status: {metadata.get('status')}")
            raise HTTPException(status_code=400, detail=f"Document not ready yet. Status: {metadata.get('status')}")
        
        # Step 3: Get summary path
        summary_path = metadata.get("summary_path")
        if not summary_path:
            print(f"[ERROR] No summary_path in metadata")
            raise HTTPException(status_code=404, detail="Document file path not found in metadata")
        
        print(f"[DEBUG] Summary path from metadata: {summary_path}")
        
        # Step 4: Check if file exists
        if not os.path.exists(summary_path):
            print(f"[ERROR] Summary file does not exist at: {summary_path}")
            # List directory contents for debugging
            summary_dir = os.path.dirname(summary_path)
            if os.path.exists(summary_dir):
                files = os.listdir(summary_dir)
                print(f"[DEBUG] Files in directory {summary_dir}: {files}")
            else:
                print(f"[ERROR] Summary directory does not exist: {summary_dir}")
            raise HTTPException(status_code=404, detail=f"Document file not found at: {summary_path}")
        
        print(f"[DEBUG] Summary file exists. Size: {os.path.getsize(summary_path)} bytes")
        
        # Step 5: Test python-docx import
        try:
            from docx import Document
            from docx.shared import Inches
            print(f"[DEBUG] python-docx imported successfully")
        except ImportError as e:
            print(f"[ERROR] Failed to import python-docx: {e}")
            raise HTTPException(status_code=500, detail=f"python-docx library not available: {e}")
        
        # Step 6: Create Word document with proper formatting
        try:
            doc = Document()
            # No title header - start directly with content
            print(f"[DEBUG] Word document created")
        except Exception as e:
            print(f"[ERROR] Failed to create Word document: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to create Word document: {e}")
        
        # Step 7: Read and parse summary content
        try:
            with open(summary_path, 'r', encoding='utf-8') as f:
                content = f.read()
            print(f"[DEBUG] Summary content read. Length: {len(content)} characters")
            
            # Parse and format the markdown content
            _add_formatted_content_to_doc(doc, content)
            print(f"[DEBUG] Formatted content added to Word document")
            
        except Exception as e:
            print(f"[ERROR] Failed to read or format summary file: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to read summary file: {e}")
        
        # Step 9: Save Word document
        try:
            word_path = summary_path.replace('.txt', '.docx')
            print(f"[DEBUG] Saving Word document to: {word_path}")
            
            # Check if directory is writable
            word_dir = os.path.dirname(word_path)
            if not os.access(word_dir, os.W_OK):
                print(f"[ERROR] Directory not writable: {word_dir}")
                raise HTTPException(status_code=500, detail=f"Directory not writable: {word_dir}")
            
            doc.save(word_path)
            print(f"[DEBUG] Word document saved successfully. Size: {os.path.getsize(word_path)} bytes")
        except Exception as e:
            print(f"[ERROR] Failed to save Word document: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to save Word document: {e}")
        
        # Step 10: Return file with forced download
        try:
            filename = f"{metadata.get('filename', 'summary')}.docx"
            print(f"[DEBUG] Returning Word document as FileResponse: {filename}")
            return FileResponse(
                word_path, 
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=filename,
                headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
            )
        except Exception as e:
            print(f"[ERROR] Failed to return FileResponse: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to return file: {e}")
            
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print(f"[ERROR] Unexpected error in export_document_word: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Unexpected error: {e}")

@app.post("/transcript/{audio_id}/edit")
async def edit_transcript(audio_id: str, request: Request):
    data = await request.json()
    transcript = data.get("transcript")
    if not transcript:
        return {"status": "error", "detail": "No transcript provided"}
    # Find the transcript path from metadata
    import os
    import pandas as pd
    metadata_dir = os.path.join("..", "data", audio_id, "metadata")
    metadata_files = [f for f in os.listdir(metadata_dir) if f.endswith('.json')]
    if not metadata_files:
        return {"status": "error", "detail": "Metadata not found"}
    
    metadata_path = os.path.join(metadata_dir, metadata_files[0])
    with open(metadata_path, 'r', encoding='utf-8') as f:
        metadata = json.load(f)
    
    transcript_path = metadata.get("transcript_path")
    if not transcript_path:
        return {"status": "error", "detail": "Transcript path not found in metadata"}
    
    # Save updated transcript
    df = pd.DataFrame(transcript)
    df.to_csv(transcript_path, index=False, encoding='utf-8')
    
    return {"status": "success", "message": "Transcript updated successfully"}

@app.post("/summary/{audio_id}/edit")
async def edit_summary(audio_id: str, request: Request):
    data = await request.json()
    summary = data.get("summary")
    apply_formatting = data.get("apply_formatting", False)
    
    if not summary:
        return {"status": "error", "detail": "No summary provided"}
    
    # Find the summary path from metadata
    metadata_dir = os.path.join("..", "data", audio_id, "metadata")
    metadata_files = [f for f in os.listdir(metadata_dir) if f.endswith('.json')]
    if not metadata_files:
        return {"status": "error", "detail": "Metadata not found"}
    
    metadata_path = os.path.join(metadata_dir, metadata_files[0])
    with open(metadata_path, 'r', encoding='utf-8') as f:
        metadata = json.load(f)
    
    summary_path = metadata.get("summary_path")
    if not summary_path:
        return {"status": "error", "detail": "Summary path not found in metadata"}
    
    # Apply formatting if requested
    final_summary = summary
    if apply_formatting:
        try:
            # Import the formatting function
            import sys
            import os
            sys.path.append(os.path.dirname(os.path.abspath(__file__)))
            from summarize_csv import format_content_with_agent
            
            final_summary = format_content_with_agent(summary)
        except Exception as e:
            print(f"[WARNING] Formatting failed: {e}")
            # Use original summary if formatting fails
            final_summary = summary
    
    # Save updated summary
    with open(summary_path, 'w', encoding='utf-8') as f:
        f.write(final_summary)
    
    return {
        "status": "success", 
        "message": "Summary updated successfully",
        "formatted_content": final_summary
    }

@app.get("/dashboard")
async def get_dashboard():
    """Get dashboard data - all audio metadata"""
    audios = get_all_audio_metadata()
    return {
        "total_audios": len(audios),
        "audios": audios
    }

@app.delete("/audio/{audio_id}")
async def delete_audio(audio_id: str):
    """Delete audio and all associated files"""
    import shutil
    
    # Path to audio directory
    audio_dir = os.path.join("..", "data", audio_id)
    
    if not os.path.exists(audio_dir):
        raise HTTPException(status_code=404, detail="Audio not found")
    
    try:
        # Remove entire audio directory
        shutil.rmtree(audio_dir)
        return {"status": "success", "message": f"Audio {audio_id} deleted successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting audio: {str(e)}")

# ============================================================================
# PROMPT MANAGEMENT ENDPOINTS
# ============================================================================

@app.get("/prompts")
async def get_prompts():
    """Get list of all available prompts"""
    try:
        prompts = list_prompts()
        return {
            "prompts": prompts,
            "total": len(prompts)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error loading prompts: {str(e)}")

@app.get("/prompts/{prompt_name}")
async def get_prompt_content(prompt_name: str):
    """Get content of a specific prompt"""
    try:
        prompt_content = get_prompt_manager().get_prompt(prompt_name)
        if not prompt_content:
            raise HTTPException(status_code=404, detail=f"Prompt '{prompt_name}' not found")
        
        return {
            "prompt_name": prompt_name,
            "content": prompt_content
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error loading prompt: {str(e)}")

@app.post("/prompts")
async def create_prompt(request: Request):
    """Create a new prompt"""
    try:
        data = await request.json()
        prompt_name = data.get("prompt_name")
        content = data.get("content")
        
        if not prompt_name or not content:
            raise HTTPException(status_code=400, detail="prompt_name and content are required")
        
        success = get_prompt_manager().add_prompt(prompt_name, content)
        if not success:
            raise HTTPException(status_code=500, detail="Failed to create prompt")
        
        return {
            "status": "success",
            "message": f"Prompt '{prompt_name}' created successfully",
            "prompt_name": prompt_name
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating prompt: {str(e)}")

@app.put("/prompts/{prompt_name}")
async def update_prompt(prompt_name: str, request: Request):
    """Update an existing prompt"""
    try:
        data = await request.json()
        content = data.get("content")
        
        if not content:
            raise HTTPException(status_code=400, detail="content is required")
        
        success = get_prompt_manager().update_prompt(prompt_name, content)
        if not success:
            raise HTTPException(status_code=404, detail=f"Prompt '{prompt_name}' not found")
        
        return {
            "status": "success",
            "message": f"Prompt '{prompt_name}' updated successfully",
            "prompt_name": prompt_name
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating prompt: {str(e)}")

@app.delete("/prompts/{prompt_name}")
async def delete_prompt(prompt_name: str):
    """Delete a prompt"""
    try:
        success = get_prompt_manager().delete_prompt(prompt_name)
        if not success:
            raise HTTPException(status_code=404, detail=f"Prompt '{prompt_name}' not found")
        
        return {
            "status": "success",
            "message": f"Prompt '{prompt_name}' deleted successfully"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting prompt: {str(e)}")

@app.post("/prompts/reload")
async def reload_prompts_endpoint():
    """Reload all prompts from files"""
    try:
        reload_prompts()
        return {
            "status": "success",
            "message": "Prompts reloaded successfully",
            "total_prompts": len(list_prompts())
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reloading prompts: {str(e)}")

@app.post("/generate-summary/{audio_id}")
async def generate_summary(audio_id: str, request: Request):
    data = await request.json()
    summary_type = data.get("summary_type", "general")
    prompt = data.get("prompt", "")
    instructions = data.get("instructions", "")
    
    # Use prompt manager to get the prompt
    if not prompt:
        # Try to get prompt from prompt manager
        prompt_content = get_prompt_manager().get_prompt(summary_type)
        if prompt_content:
            prompt = prompt_content
        else:
            # Fallback to hardcoded prompts
            PROMPT_TEMPLATES = {
                "general": "You are an expert business analyst and technical writer.\nBased on the meeting transcript I will provide, identify and extract all distinct business requirements discussed by the stakeholders.\n...",
                "fsd": "You are an expert functional specification document (FSD) writer.\nBased on the transcript, generate a detailed FSD.\n...",
            }
            prompt = PROMPT_TEMPLATES.get(summary_type, PROMPT_TEMPLATES["general"])
    
    # Format the prompt with transcript placeholder
    if "{transcript}" in prompt:
        # This will be replaced by the actual transcript in the Celery task
        pass
    
    if instructions:
        prompt = f"{prompt}\n\nAdditional instructions from user: {instructions}"
    
    metadata = get_audio_metadata(audio_id)
    if not metadata:
        raise HTTPException(status_code=404, detail="Audio not found")
    if metadata.get("status") not in ["transcribed", "summary_generated"]:
        raise HTTPException(status_code=400, detail="Transcript not ready yet")
    
    # Get metadata path to update status
    metadata_dir = os.path.join("..", "data", audio_id, "metadata")
    metadata_files = [f for f in os.listdir(metadata_dir) if f.endswith('.json')]
    if not metadata_files:
        raise HTTPException(status_code=404, detail="Metadata not found")
    
    metadata_path = os.path.join(metadata_dir, metadata_files[0])
    
    # Store prompt information in metadata
    metadata["status"] = "summary_regenerating"
    metadata["regeneration_started_at"] = datetime.now().isoformat()
    metadata["summary_started_at"] = datetime.now().isoformat()
    metadata["summary_config"] = {
        "summary_type": summary_type,
        "prompt_used": prompt,
        "instructions": instructions
    }
    update_metadata(metadata_path, metadata)
    
    # Start background summary generation with custom prompt
    task = generate_summary_task.delay(audio_id=audio_id, prompt=prompt)
    return {
        "audio_id": audio_id,
        "task_id": task.id,
        "status": "summary_regenerating",
        "message": "Summary regeneration started"
    }

@app.post("/cancel-task/{audio_id}")
async def cancel_task(audio_id: str):
    """Cancel a running task for the given audio ID"""
    try:
        from celery_worker import celery_app
        
        # Get metadata to find task ID
        metadata = get_audio_metadata(audio_id)
        if not metadata:
            raise HTTPException(status_code=404, detail="Audio not found")
        
        # Check if there's a task ID in metadata
        task_id = metadata.get("task_id")
        if not task_id:
            raise HTTPException(status_code=400, detail="No active task found for this audio")
        
        # Try to revoke the task (thread pool doesn't support termination)
        print(f"[DEBUG] Attempting to cancel task {task_id} for audio {audio_id}")
        try:
            # With thread pool, we can only revoke without termination
            result = celery_app.control.revoke(task_id, terminate=False)
            print(f"[DEBUG] Task {task_id} revoked successfully: {result}")
            
            # Note: Thread pool doesn't support immediate termination
            # The task will continue running but won't be picked up again
            print(f"[INFO] Task {task_id} revoked. It may continue running until completion.")
            
        except Exception as e:
            print(f"[ERROR] Failed to revoke task {task_id}: {e}")
            import traceback
            traceback.print_exc()
            # Continue anyway to update metadata
        
        # Update metadata to reflect cancellation
        metadata_dir = os.path.join("..", "data", audio_id, "metadata")
        metadata_files = [f for f in os.listdir(metadata_dir) if f.endswith('.json')]
        if metadata_files:
            metadata_path = os.path.join(metadata_dir, metadata_files[0])
            metadata["status"] = "cancelled"
            metadata["cancelled_at"] = datetime.now().isoformat()
            update_metadata(metadata_path, metadata)
        
        return {
            "status": "success",
            "message": f"Task cancelled successfully",
            "audio_id": audio_id
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error cancelling task: {str(e)}")

@app.get("/timing-estimate")
async def get_timing_estimate(
    audio_duration_minutes: float = 10,
    diarizer: str = "huggingface",
    speedup: float = 1.0,
    chunk_mode: bool = False,
    chunk_duration: int = 10,
    summary_type: str = "general",
    transcript_length_chars: int = 10000
):
    """Get timing estimates for audio processing and summary generation"""
    try:
        # Get audio processing estimate
        audio_estimate, audio_confidence = timing_model.estimate_audio_processing_time(
            audio_duration_minutes=audio_duration_minutes,
            diarizer=diarizer,
            speedup=speedup,
            chunk_mode=chunk_mode,
            chunk_duration=chunk_duration
        )
        
        # Get summary generation estimate
        summary_estimate, summary_confidence = timing_model.estimate_summary_generation_time(
            transcript_length_chars=transcript_length_chars,
            summary_type=summary_type
        )
        
        return {
            "audio_processing": {
                "estimated_seconds": audio_estimate,
                "confidence": audio_confidence,
                "estimated_minutes": round(audio_estimate / 60, 1)
            },
            "summary_generation": {
                "estimated_seconds": summary_estimate,
                "confidence": summary_confidence,
                "estimated_minutes": round(summary_estimate / 60, 1)
            },
            "timing_stats": timing_model.get_timing_stats()
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting timing estimate: {str(e)}")

@app.get("/download-audio/{audio_id}")
async def download_audio(audio_id: str):
    """Download the original audio file"""
    try:
        print(f"[DEBUG] Audio download requested for audio_id: {audio_id}")
        
        # Get metadata to get filename
        metadata = get_audio_metadata(audio_id)
        if not metadata:
            print(f"[ERROR] No metadata found for audio_id: {audio_id}")
            raise HTTPException(status_code=404, detail="Audio not found")
        
        # Find the audio file
        try:
            from utils import find_input_audio
            audio_path = find_input_audio(audio_id)
            print(f"[DEBUG] Audio file found at: {audio_path}")
        except FileNotFoundError as e:
            print(f"[ERROR] Audio file not found: {e}")
            raise HTTPException(status_code=404, detail="Audio file not found")
        
        # Check if file exists
        if not os.path.exists(audio_path):
            print(f"[ERROR] Audio file does not exist at: {audio_path}")
            raise HTTPException(status_code=404, detail=f"Audio file not found at: {audio_path}")
        
        # Get original filename from metadata
        original_filename = metadata.get('filename', 'audio.wav')
        print(f"[DEBUG] Original filename: {original_filename}")
        
        # Return file with forced download
        return FileResponse(
            audio_path,
            media_type="application/octet-stream",  # Force download
            filename=original_filename,
            headers={"Content-Disposition": f"attachment; filename=\"{original_filename}\""}
        )
        
    except Exception as e:
        print(f"[ERROR] Unexpected error in download_audio: {e}")
        raise HTTPException(status_code=500, detail=f"Error downloading audio: {str(e)}")
